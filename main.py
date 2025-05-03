import os
import time
import json
import hashlib
import sys
import asyncio
from pathlib import Path
from typing import List, Optional, Set, Dict, Any, Tuple

import typer
from docx import Document
from docx.enum.text import WD_BREAK
import openai
from openai import OpenAI
from openai import AsyncOpenAI
from loguru import logger
from tqdm import tqdm
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Create Typer app
app = typer.Typer(help="Translate a Word document using OpenAI API")

# Constants
DEFAULT_MODEL = "gpt-4-turbo"
MAX_RETRIES = 5
INITIAL_RETRY_DELAY = 1
TEMPERATURE = 0.3
DEFAULT_CACHE_DIR = Path.home() / ".cache/.docx_translator"
DEFAULT_MAX_CONCURRENT = 5
DEFAULT_TARGET_LANGUAGE = os.environ.get("TARGET_LANGUAGE", "Spanish")

# Configure logging with loguru
logger.remove()  # Remove default handler
logger.add(
    sys.stderr,
    format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
    level="INFO",
    colorize=True,
)
logger.add(
    "docx_translator.log",
    rotation="10 MB",
    retention="1 week",
    level="DEBUG",
    compression="zip",
)

# Create a dictionary to track streaming progress
translation_progress = {}
total_tokens_received = 0
total_cached_tokens = 0  # Track total tokens from cache


def setup_openai_client(
    api_key: Optional[str] = None, base_url: Optional[str] = None
) -> OpenAI:
    """Set up and return OpenAI client

    Args:
        api_key: OpenAI API key (optional if environment variable is set)
        base_url: OpenAI API base URL for custom endpoints (optional)

    Returns:
        OpenAI client instance

    Raises:
        ValueError: If API key is not provided and not found in environment
    """
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        raise ValueError(
            "OpenAI API key is required. Set it via --api_key or OPENAI_API_KEY environment variable."
        )

    client_kwargs = {"api_key": key}

    # Add base_url if provided
    if base_url:
        client_kwargs["base_url"] = base_url
        logger.info(f"Using custom OpenAI API base URL: {base_url}")

    return OpenAI(**client_kwargs)


def setup_async_openai_client(
    api_key: Optional[str] = None, base_url: Optional[str] = None
) -> AsyncOpenAI:
    """Set up and return async OpenAI client

    Args:
        api_key: OpenAI API key (optional if environment variable is set)
        base_url: OpenAI API base URL for custom endpoints (optional)

    Returns:
        AsyncOpenAI client instance

    Raises:
        ValueError: If API key is not provided and not found in environment
    """
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        raise ValueError(
            "OpenAI API key is required. Set it via --api_key or OPENAI_API_KEY environment variable."
        )

    client_kwargs = {"api_key": key}

    # Add base_url if provided
    if base_url:
        client_kwargs["base_url"] = base_url
        logger.info(f"Using custom OpenAI API base URL: {base_url}")

    return AsyncOpenAI(**client_kwargs)


class TranslationCache:
    """Cache for translations to avoid redundant API calls"""

    def __init__(
        self,
        target_language: str,
        model: str = DEFAULT_MODEL,
        cache_dir: Path = DEFAULT_CACHE_DIR,
    ):
        """Initialize the translation cache

        Args:
            target_language: The target language for translations
            model: The model used for translations
            cache_dir: Directory to store cache files
        """
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True, parents=True)

        # Create a cache file specific to this language and model
        cache_filename = f"cache_{target_language.lower().replace(' ', '_')}_{model.replace('-', '_')}.json"
        self.cache_file = self.cache_dir / cache_filename

        # Load existing cache if available
        self.cache: Dict[str, Dict[str, Any]] = {}
        if self.cache_file.exists():
            try:
                with open(self.cache_file, "r", encoding="utf-8") as f:
                    self.cache = json.load(f)
                logger.info(
                    f"Loaded {len(self.cache)} cached translations for {target_language}"
                )
            except Exception as e:
                logger.warning(f"Failed to load translation cache: {e}")
                # Start with an empty cache if loading fails
                self.cache = {}

    def get_cache_key(self, text: str) -> str:
        """Generate a unique key for the text

        Args:
            text: The text to be cached

        Returns:
            A unique hash for the text
        """
        # Use MD5 hash of the text as the key
        return hashlib.md5(text.encode("utf-8")).hexdigest()

    def get(self, text: str) -> Optional[Dict[str, Any]]:
        """Get a translation from the cache

        Args:
            text: The text to get translation for

        Returns:
            Dictionary with translated text and token usage or None if not found
        """
        key = self.get_cache_key(text)
        return self.cache.get(key)

    def set(
        self, text: str, translation: str, token_usage: Optional[Dict[str, int]] = None
    ) -> None:
        """Store a translation in the cache with token usage information

        Args:
            text: The original text
            translation: The translated text
            token_usage: Dictionary with token usage information
        """
        key = self.get_cache_key(text)

        # Store both the translation and token usage
        self.cache[key] = {
            "translation": translation,
            "token_usage": token_usage or {},
        }

        # Periodically save the cache to disk (every 10 new entries)
        if len(self.cache) % 10 == 0:
            self.save()

    def save(self) -> None:
        """Save the cache to disk"""
        try:
            with open(self.cache_file, "w", encoding="utf-8") as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
            logger.debug(f"Cache saved with {len(self.cache)} entries")
        except Exception as e:
            logger.warning(f"Failed to save translation cache: {e}")


async def stream_async_translation(
    text: str,
    target_language: str,
    client: AsyncOpenAI,
    cache: Optional[TranslationCache] = None,
    progress_bar: Optional[tqdm] = None,
    task_id: Optional[str] = None,
    cancellation_check=None,
) -> Tuple[str, str]:
    """Stream translation using OpenAI API asynchronously with token counting

    Args:
        text: Text to translate
        target_language: Target language for translation
        client: AsyncOpenAI client instance
        cache: Optional translation cache
        progress_bar: tqdm progress bar to update
        task_id: Identifier for this translation task
        cancellation_check: Optional function that returns True if the task should be cancelled

    Returns:
        Tuple of (original text, translated text)
    """
    global total_tokens_received, total_cached_tokens

    # Skip empty text
    if not text.strip():
        return text, ""

    # Check for cancellation before starting
    if cancellation_check and cancellation_check():
        # Log without raising exception directly - let it propagate through CancelledError
        # which is handled more gracefully
        logger.warning(f"Task {task_id} cancelled before starting translation")
        raise asyncio.CancelledError("Translation cancelled by user")

    # Check cache first if available
    if cache:
        cached_data = cache.get(text)
        if cached_data:
            logger.debug(f"Using cached translation for: {text[:40]}...")

            # Update token counter with cached token usage if available
            if (
                cached_data.get("token_usage")
                and "completion_tokens" in cached_data["token_usage"]
            ):
                completion_tokens = cached_data["token_usage"]["completion_tokens"]
                total_cached_tokens += completion_tokens
                if progress_bar:
                    progress_bar.set_description(
                        f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                    )

            return text, cached_data["translation"]

    retry_delay = INITIAL_RETRY_DELAY
    network_retry_count = 0
    MAX_NETWORK_RETRIES = 3

    # Initialize task in translation progress tracking
    if task_id:
        translation_progress[task_id] = 0

    for attempt in range(MAX_RETRIES):
        # Check for cancellation before each attempt
        if cancellation_check and cancellation_check():
            logger.warning(f"Task {task_id} cancelled before attempt {attempt+1}")
            raise asyncio.CancelledError("Translation cancelled by user")

        try:
            # Initialize the collection of streamed content
            collected_content = ""
            token_usage = None

            # Call the OpenAI API asynchronously with streaming
            stream = await client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the following text to {target_language}. Preserve formatting and maintain the same tone and style.",
                    },
                    {"role": "user", "content": text},
                ],
                temperature=TEMPERATURE,
                stream=True,
            )

            # Process the stream
            async for chunk in stream:
                # Check for cancellation during streaming
                if cancellation_check and cancellation_check():
                    logger.warning(f"Task {task_id} cancelled during streaming")
                    raise asyncio.CancelledError("Translation cancelled by user")

                if hasattr(chunk, "choices") and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    if hasattr(delta, "content") and delta.content:
                        # Append content to our collection
                        collected_content += delta.content

                        # Update the progress tracking
                        if task_id:
                            translation_progress[task_id] = len(collected_content)

                # Check if we have usage information (only in the final chunk)
                if hasattr(chunk, "usage") and chunk.usage is not None:
                    token_usage = {
                        "completion_tokens": chunk.usage.completion_tokens,
                        "prompt_tokens": chunk.usage.prompt_tokens,
                        "total_tokens": chunk.usage.total_tokens,
                    }

                    # Update the global token count and progress bar
                    if progress_bar and "completion_tokens" in token_usage:
                        total_tokens_received += token_usage["completion_tokens"]
                        progress_bar.set_description(
                            f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                        )

            # Check for cancellation after streaming but before post-processing
            if cancellation_check and cancellation_check():
                logger.warning(f"Task {task_id} cancelled after streaming")
                raise asyncio.CancelledError("Translation cancelled by user")

            # For non-streaming response, update token count if we have it
            if not token_usage:
                # Make a non-streaming call to get token usage
                non_stream_response = await client.chat.completions.create(
                    model=DEFAULT_MODEL,
                    messages=[
                        {
                            "role": "system",
                            "content": f"You are a professional translator. Translate the following text to {target_language}. Preserve formatting and maintain the same tone and style.",
                        },
                        {"role": "user", "content": text},
                    ],
                    temperature=TEMPERATURE,
                    stream=False,
                )

                if (
                    hasattr(non_stream_response, "usage")
                    and non_stream_response.usage is not None
                ):
                    token_usage = {
                        "completion_tokens": non_stream_response.usage.completion_tokens,
                        "prompt_tokens": non_stream_response.usage.prompt_tokens,
                        "total_tokens": non_stream_response.usage.total_tokens,
                    }

                    # Update the global token count and progress bar
                    if progress_bar and "completion_tokens" in token_usage:
                        total_tokens_received += token_usage["completion_tokens"]
                        progress_bar.set_description(
                            f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                        )

            # Reset network retry counter on successful request
            network_retry_count = 0

            translation = collected_content.strip()

            # Store in cache if available and we have a complete translation
            if cache and translation:
                cache.set(text, translation, token_usage)

            return text, translation
        except openai.RateLimitError:
            if attempt < MAX_RETRIES - 1:
                logger.warning(
                    f"Rate limit exceeded. Retrying in {retry_delay} seconds..."
                )
                await asyncio.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                logger.error(
                    f"Failed after {MAX_RETRIES} attempts due to rate limiting"
                )
                return (
                    text,
                    f"[Translation Error: Rate limit exceeded after {MAX_RETRIES} attempts]",
                )
        except (
            openai.APIConnectionError,
            openai.AuthenticationError,
            openai.APITimeoutError,
            asyncio.TimeoutError,
        ) as e:
            # Handle network-related errors with a separate retry counter
            if network_retry_count < MAX_NETWORK_RETRIES:
                network_retry_count += 1
                retry_delay_network = 1 * network_retry_count  # Linear backoff
                logger.warning(
                    f"Network error: {str(e)}. Network retry {network_retry_count}/{MAX_NETWORK_RETRIES} in {retry_delay_network} seconds..."
                )
                await asyncio.sleep(retry_delay_network)
                # Don't increment the main attempt counter for network errors
                attempt -= 1
            else:
                logger.error(
                    f"Failed after {MAX_NETWORK_RETRIES} network retry attempts: {str(e)}"
                )
                return (
                    text,
                    f"[Translation Error: Network error after {MAX_NETWORK_RETRIES} retry attempts]",
                )
        except asyncio.CancelledError:
            # Handle cancellation by properly ending the task
            logger.warning(f"Translation task {task_id} canceled")
            return text, f"[Translation Canceled]"
        except Exception as e:
            logger.exception(f"Error translating text")
            return text, f"[Translation Error: {str(e)}]"


def stream_translation(
    text: str,
    target_language: str,
    client: OpenAI,
    cache: Optional[TranslationCache] = None,
    progress_bar: Optional[tqdm] = None,
    task_id: Optional[str] = None,
) -> str:
    """Translate text using OpenAI API streaming with token counting

    Args:
        text: Text to translate
        target_language: Target language for translation
        client: OpenAI client instance
        cache: Optional translation cache
        progress_bar: tqdm progress bar to update
        task_id: Identifier for this translation task

    Returns:
        Translated text or error message
    """
    global total_tokens_received, total_cached_tokens

    # Skip empty text
    if not text.strip():
        return ""

    # Check cache first if available
    if cache:
        cached_data = cache.get(text)
        if cached_data:
            logger.debug(f"Using cached translation for: {text[:40]}...")

            # Update token counter with cached token usage if available
            if (
                cached_data.get("token_usage")
                and "completion_tokens" in cached_data["token_usage"]
            ):
                completion_tokens = cached_data["token_usage"]["completion_tokens"]
                total_cached_tokens += completion_tokens
                if progress_bar:
                    progress_bar.set_description(
                        f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                    )

            return cached_data["translation"]

    retry_delay = INITIAL_RETRY_DELAY
    network_retry_count = 0
    MAX_NETWORK_RETRIES = 3

    # Initialize task in translation progress tracking
    if task_id:
        translation_progress[task_id] = 0

    for attempt in range(MAX_RETRIES):
        try:
            # Initialize the collection of streamed content
            collected_content = ""
            token_usage = None

            # Call the OpenAI API with streaming
            stream = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the following text to {target_language}. Preserve formatting and maintain the same tone and style.",
                    },
                    {"role": "user", "content": text},
                ],
                temperature=TEMPERATURE,
                stream=True,
            )

            # Process the stream
            for chunk in stream:
                if hasattr(chunk, "choices") and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    if hasattr(delta, "content") and delta.content:
                        # Append content to our collection
                        collected_content += delta.content

                        # Update the progress
                        if task_id:
                            translation_progress[task_id] = len(collected_content)

                # Check if we have usage information (only in the final chunk)
                if hasattr(chunk, "usage") and chunk.usage is not None:
                    token_usage = {
                        "completion_tokens": chunk.usage.completion_tokens,
                        "prompt_tokens": chunk.usage.prompt_tokens,
                        "total_tokens": chunk.usage.total_tokens,
                    }

                    # Update the global token count and progress bar
                    if progress_bar and "completion_tokens" in token_usage:
                        total_tokens_received += token_usage["completion_tokens"]
                        progress_bar.set_description(
                            f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                        )

            # For non-streaming response, update token count if we have it
            if not token_usage:
                # Make a non-streaming call to get token usage
                non_stream_response = client.chat.completions.create(
                    model=DEFAULT_MODEL,
                    messages=[
                        {
                            "role": "system",
                            "content": f"You are a professional translator. Translate the following text to {target_language}. Preserve formatting and maintain the same tone and style.",
                        },
                        {"role": "user", "content": text},
                    ],
                    temperature=TEMPERATURE,
                    stream=False,
                )

                if (
                    hasattr(non_stream_response, "usage")
                    and non_stream_response.usage is not None
                ):
                    token_usage = {
                        "completion_tokens": non_stream_response.usage.completion_tokens,
                        "prompt_tokens": non_stream_response.usage.prompt_tokens,
                        "total_tokens": non_stream_response.usage.total_tokens,
                    }

                    # Update the global token count and progress bar
                    if progress_bar and "completion_tokens" in token_usage:
                        total_tokens_received += token_usage["completion_tokens"]
                        progress_bar.set_description(
                            f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                        )

            # Reset network retry counter on successful request
            network_retry_count = 0

            translation = collected_content.strip()

            # Store in cache if available and we have a complete translation
            if cache and translation:
                cache.set(text, translation, token_usage)

            return translation
        except openai.RateLimitError:
            if attempt < MAX_RETRIES - 1:
                logger.warning(
                    f"Rate limit exceeded. Retrying in {retry_delay} seconds..."
                )
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                logger.error(
                    f"Failed after {MAX_RETRIES} attempts due to rate limiting"
                )
                return f"[Translation Error: Rate limit exceeded after {MAX_RETRIES} attempts]"
        except (openai.APIConnectionError, openai.APITimeoutError) as e:
            # Handle network-related errors with a separate retry counter
            if network_retry_count < MAX_NETWORK_RETRIES:
                network_retry_count += 1
                retry_delay_network = 1 * network_retry_count  # Linear backoff
                logger.warning(
                    f"Network error: {str(e)}. Network retry {network_retry_count}/{MAX_NETWORK_RETRIES} in {retry_delay_network} seconds..."
                )
                time.sleep(retry_delay_network)
                # Don't increment the main attempt counter for network errors
                attempt -= 1
            else:
                logger.error(
                    f"Failed after {MAX_NETWORK_RETRIES} network retry attempts: {str(e)}"
                )
                return f"[Translation Error: Network error after {MAX_NETWORK_RETRIES} retry attempts]"
        except Exception as e:
            logger.exception(f"Error translating text")
            return f"[Translation Error: {str(e)}]"


def add_translation_to_paragraph(paragraph, translated_text: str) -> None:
    """Add translation to a paragraph with appropriate formatting

    Args:
        paragraph: Document paragraph to modify
        translated_text: Text to add as translation
    """
    # Preserve the existing paragraph content and styling
    # Simply append a line break and the translated text at the end

    # Find the last run or create one if none exists
    if len(paragraph.runs) > 0:
        last_run = paragraph.runs[-1]
        # Only add a line break if the text doesn't already end with one
        if not last_run.text.endswith("\n"):
            last_run.add_break(WD_BREAK.LINE)
    else:
        # If no runs exist (rare case), use the paragraph's text property
        last_run = paragraph.add_run(paragraph.text)
        last_run.add_break(WD_BREAK.LINE)

    # Add a separator line
    separator_run = paragraph.add_run("---------------------")
    separator_run.add_break(WD_BREAK.LINE)

    # Add the translated text as a new run
    translation_run = paragraph.add_run(translated_text)

    # If this was a heading, maintain formatting for translation too
    if paragraph.style.name.startswith("Heading"):
        translation_run.bold = True


def count_translatable_elements(doc: Document, target_styles_set: Set[str]) -> int:
    """Count elements that need translation based on their style

    Args:
        doc: Document to analyze
        target_styles_set: Set of style names to count

    Returns:
        Number of elements to translate
    """
    count = 0

    # Count paragraphs
    for para in doc.paragraphs:
        if para.style.name in target_styles_set and para.text.strip():
            count += 1

    # Count elements in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style.name in target_styles_set and para.text.strip():
                        count += 1

    return count


async def process_document_parallel(
    input_file: Path,
    output_file: Path,
    target_language: str,
    target_styles: List[str],
    openai_client: OpenAI,
    use_cache: bool = True,
    cache_dir: Path = DEFAULT_CACHE_DIR,
    max_concurrent: int = DEFAULT_MAX_CONCURRENT,
    progress_callback=None,
    cancellation_check=None,
) -> None:
    """Process the document with parallel translation requests

    Args:
        input_file: Path to input DOCX file
        output_file: Path to output DOCX file
        target_language: Target language for translation
        target_styles: List of paragraph styles to translate
        openai_client: OpenAI client instance
        use_cache: Whether to use translation caching
        cache_dir: Directory to store cache files
        max_concurrent: Maximum number of concurrent translation requests
        progress_callback: Optional callback function to update progress (signature: current, total, description)
        cancellation_check: Optional function that returns True if the process should be cancelled
    """
    logger.info(f"Opening document: {input_file}")
    doc = Document(input_file)

    # Convert target_styles to set for faster lookups
    target_styles_set = set(target_styles)

    # Initialize async OpenAI client using the same credentials
    async_client = setup_async_openai_client(
        openai_client.api_key, getattr(openai_client, "base_url", None)
    )

    # Initialize translation cache if enabled
    cache = (
        TranslationCache(target_language, DEFAULT_MODEL, cache_dir)
        if use_cache
        else None
    )
    if cache:
        logger.info("Translation caching is enabled")

    # Collect all translatable paragraphs
    all_paragraphs = []
    paragraphs_data = []

    # Process main paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and para.style.name in target_styles_set:
            paragraphs_data.append((para, text))
            all_paragraphs.append(para)

    # Process table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and para.style.name in target_styles_set:
                        paragraphs_data.append((para, text))
                        all_paragraphs.append(para)

    total_elements = len(paragraphs_data)
    logger.info(f"Translating document to {target_language}")
    logger.info(f"Targeting styles: {', '.join(target_styles)}")
    logger.info(f"Elements to translate: {total_elements}")
    logger.info(
        f"Using parallel processing with max {max_concurrent} concurrent requests"
    )

    # Reset the token counters
    global total_tokens_received, total_cached_tokens
    total_tokens_received = 0
    total_cached_tokens = 0

    # Create a progress bar
    progress_bar = tqdm(
        total=total_elements,
        desc=f"Translating (0 API + 0 cached = 0 total tokens)",
    )

    # Create a semaphore to limit concurrent requests
    semaphore = asyncio.Semaphore(max_concurrent)

    # Create tasks for unique text only
    unique_texts = {}
    unique_tasks = {}
    task_map = {}

    # Store unique texts and create mappings
    for idx, (para, text) in enumerate(paragraphs_data):
        if text in unique_texts:
            # Map this paragraph to existing task
            source_idx = unique_texts[text]
            task_map[idx] = source_idx
        else:
            # This is a new unique text
            source_idx = len(unique_tasks)
            unique_texts[text] = source_idx
            unique_tasks[source_idx] = (para, text)
            task_map[idx] = source_idx

    logger.info(
        f"Found {len(unique_tasks)} unique text elements to translate out of {total_elements} total"
    )

    cached_hits = 0
    cached_tokens = 0

    # Check cache first
    if use_cache and cache:
        for idx, (para, text) in unique_tasks.items():
            cache_result = cache.get(text)
            if cache_result and cache_result.get("translation"):
                cached_hits += 1
                if (
                    cache_result.get("token_usage")
                    and "completion_tokens" in cache_result["token_usage"]
                ):
                    cached_tokens += cache_result["token_usage"]["completion_tokens"]
                    total_cached_tokens += cache_result["token_usage"][
                        "completion_tokens"
                    ]

        logger.info(f"Found {cached_hits} translations in cache")

    # Define the async translation function with semaphore
    try:
        tasks = []  # Track all tasks for proper cleanup

        async def translate_with_semaphore(text, task_num, task_idx):
            """Translate text using semaphore to limit concurrent requests"""
            # Check for cancellation before acquiring semaphore
            if cancellation_check and cancellation_check():
                logger.warning(f"Task {task_num} cancelled before starting")
                raise asyncio.CancelledError("Translation cancelled by user")

            # Use the semaphore to limit concurrent requests
            async with semaphore:
                # Check for cancellation again after acquiring semaphore
                if cancellation_check and cancellation_check():
                    logger.warning(
                        f"Task {task_num} cancelled after acquiring semaphore"
                    )
                    raise asyncio.CancelledError("Translation cancelled by user")

                if progress_callback:
                    # Get progress directly from tqdm for consistency
                    current_progress = progress_bar.n
                    eta_seconds = progress_bar.format_dict.get("eta", None)
                    eta_str = (
                        f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                        if eta_seconds is not None
                        else ""
                    )
                    progress_callback(
                        current_progress,
                        total_elements,
                        f"Preparing task {task_num}/{len(unique_tasks)}{eta_str}",
                    )

                # First check cache
                if use_cache and cache:
                    cache_result = cache.get(text)
                    if cache_result and cache_result.get("translation"):
                        # Explicitly update progress bar count here
                        progress_bar.update(1)
                        progress_bar.set_description(
                            f"Translating ({total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total tokens)"
                        )

                        if progress_callback:
                            # Get progress directly from tqdm for consistency
                            current_progress = progress_bar.n
                            eta_seconds = progress_bar.format_dict.get("eta", None)
                            eta_str = (
                                f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                                if eta_seconds is not None
                                else ""
                            )
                            progress_callback(
                                current_progress,
                                total_elements,
                                f"Task {task_num}/{len(unique_tasks)} (cached): {text[:30]}...{eta_str}",
                            )

                        return text, cache_result["translation"]

                # Not in cache, translate
                task_id = f"Task {task_num}/{len(unique_tasks)}"
                if progress_callback:
                    # Get progress directly from tqdm for consistency
                    current_progress = progress_bar.n
                    eta_seconds = progress_bar.format_dict.get("eta", None)
                    eta_str = (
                        f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                        if eta_seconds is not None
                        else ""
                    )
                    progress_callback(
                        current_progress,
                        total_elements,
                        f"Translating {task_id}: {text[:30]}...{eta_str}",
                    )

                # Use streaming translation to get progress feedback
                original, translated = await stream_async_translation(
                    text,
                    target_language,
                    async_client,
                    cache,
                    progress_bar,
                    task_id,
                    cancellation_check,
                )

                # Explicitly update progress bar count here for non-cached translations
                progress_bar.update(1)

                # Update Streamlit progress with tqdm's current count and ETA
                if progress_callback:
                    current_progress = progress_bar.n
                    eta_seconds = progress_bar.format_dict.get("eta", None)
                    eta_str = (
                        f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                        if eta_seconds is not None
                        else ""
                    )
                    progress_callback(
                        current_progress,
                        total_elements,
                        f"Completed {task_id}: {text[:30]}...{eta_str}",
                    )

                return original, translated

        # Create and gather all translation tasks
        translation_tasks = []
        for task_idx, (para, text) in unique_tasks.items():
            # Create a coroutine with both task number (for display) and task index (for mapping)
            coroutine = translate_with_semaphore(text, task_idx + 1, task_idx)
            # Wrap it in a Task so it can be properly cancelled
            task = asyncio.create_task(coroutine)
            translation_tasks.append(task)
            tasks.append(task)  # Keep track for cleanup

        # Check for initial cancellation before starting tasks
        if cancellation_check and cancellation_check():
            raise asyncio.CancelledError(
                "Translation cancelled by user before starting"
            )

        try:
            # Execute all tasks concurrently with the semaphore limiting concurrent execution
            translations = await asyncio.gather(*translation_tasks)
        except asyncio.CancelledError:
            logger.warning("Translation tasks cancelled")
            raise  # Re-raise to be caught by the outer try block

        # Check for cancellation before applying translations
        if cancellation_check and cancellation_check():
            raise asyncio.CancelledError(
                "Translation cancelled by user after tasks completed"
            )

        # Create a dictionary mapping task_idx to translation
        translations_dict = {
            idx: trans for idx, (orig, trans) in enumerate(translations)
        }

        # Apply translations to all paragraphs
        paragraphs_processed = 0
        for idx, para in enumerate(all_paragraphs):
            try:
                # Check for cancellation periodically while applying translations
                if cancellation_check and idx % 10 == 0 and cancellation_check():
                    raise asyncio.CancelledError(
                        "Translation cancelled while applying results"
                    )

                # Get the source task index from the map
                source_idx = task_map.get(idx)
                if source_idx is None:
                    logger.warning(
                        f"No source index found for paragraph {idx}. Skipping translation."
                    )
                    # Update progress bar even for skipped items
                    progress_bar.update(1)
                    continue

                # Get the translation for that source
                translation = translations_dict.get(source_idx)
                if translation is None:
                    logger.warning(
                        f"No translation found for source index {source_idx} (paragraph {idx}). Skipping."
                    )
                    # Update progress bar even for skipped items
                    progress_bar.update(1)
                    continue

                # Add it to the paragraph
                add_translation_to_paragraph(para, translation)
                paragraphs_processed += 1

            except asyncio.CancelledError:
                raise  # Re-raise cancellation
            except Exception as e:
                logger.error(f"Error applying translation to paragraph {idx}: {str(e)}")
                logger.exception(e)
                logger.debug(f"Paragraph text: {para.text[:100]}...")
                # Update progress bar even for error items
                progress_bar.update(1)
                continue

            # Update progress if callback is provided (for overall progress visibility)
            if progress_callback and idx % 10 == 0:
                # Try to get ETA from tqdm in different ways
                eta_seconds = progress_bar.format_dict.get("eta", None)

                # If eta is None, try to calculate our own ETA based on rate
                if eta_seconds is None and progress_bar.n > 0:
                    rate = progress_bar.format_dict.get("rate", None)
                    if rate and rate > 0:
                        remaining = progress_bar.total - progress_bar.n
                        calculated_eta = remaining / rate
                        eta_seconds = calculated_eta
                        logger.debug(
                            f"Calculated our own ETA: {calculated_eta:.1f}s based on rate {rate:.2f} items/sec"
                        )

                logger.debug(f"TQDM format_dict: {progress_bar.format_dict}")
                logger.debug(f"ETA seconds (final): {eta_seconds}")

                eta_str = (
                    f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                    if eta_seconds is not None
                    else ""
                )
                logger.debug(f"Formatted ETA string: {eta_str}")
                progress_callback(
                    progress_bar.n,  # Use tqdm's progress as source of truth
                    total_elements,
                    f"Applied {paragraphs_processed} translations{eta_str}",
                )

        # Close progress bar
        progress_bar.close()

        # Log cache stats
        if use_cache:
            logger.info(f"Used {cached_hits} cached translations")
            logger.info(f"Saved approximately {cached_tokens} tokens with caching")

        # Save document
        logger.info(f"Saving translated document to: {output_file}")
        doc.save(output_file)

        # Calculate how many API calls were made
        api_calls = len(unique_tasks) - cached_hits
        logger.success(
            f"Document translation completed! Translated {paragraphs_processed} elements ({api_calls} API calls, {cached_hits} from cache) to {target_language}"
        )
        logger.info(
            f"Total tokens used: {total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total"
        )

        # Final progress update
        if progress_callback:
            progress_callback(total_elements, total_elements, "Translation completed!")

    except asyncio.CancelledError:
        # Handle cancellation by the user
        logger.warning("Translation was cancelled by the user")

        # Close progress bar
        progress_bar.close()

        # Cancel any pending tasks to prevent progress bar from continuing
        for task in tasks:
            # Handle different task types correctly
            try:
                # Check if the task is a coroutine or an asyncio Task
                if hasattr(task, "cancel"):
                    # It's a Future or Task object
                    if not task.done():
                        task.cancel()
                else:
                    # For coroutine objects, we can't cancel directly
                    # But they should be cleaned up when the event loop stops
                    pass
            except Exception as e:
                logger.warning(f"Error handling task cancellation: {str(e)}")

        # Wait for tasks to be cancelled
        try:
            await asyncio.gather(*tasks, return_exceptions=True)
        except asyncio.CancelledError:
            pass

        # Raise a specific exception that can be caught by the caller
        raise InterruptedError("Translation was cancelled by the user")
    except InterruptedError as e:
        # For user-triggered cancellation, use warning without stack trace
        logger.warning(f"Translation was interrupted: {str(e)}")
        raise
    except Exception as e:
        # Close progress bar on error
        progress_bar.close()

        # Cancel any pending tasks to prevent progress bar from continuing
        for task in tasks:
            # Handle different task types correctly
            try:
                # Check if the task is a coroutine or an asyncio Task
                if hasattr(task, "cancel"):
                    # It's a Future or Task object
                    if not task.done():
                        task.cancel()
                else:
                    # For coroutine objects, we can't cancel directly
                    # But they should be cleaned up when the event loop stops
                    pass
            except Exception as e:
                logger.warning(f"Error handling task cancellation: {str(e)}")

        # Wait for tasks to be cancelled
        try:
            await asyncio.gather(*tasks, return_exceptions=True)
        except asyncio.CancelledError:
            pass

        # Re-raise the exception to be handled by the caller
        logger.exception(e)
        logger.error(f"Error during parallel translation: {str(e)}")
        raise e


def process_document(
    input_file: Path,
    output_file: Path,
    target_language: str,
    target_styles: List[str],
    openai_client: OpenAI,
    use_cache: bool = True,
    cache_dir: Path = DEFAULT_CACHE_DIR,
    parallel: bool = True,
    max_concurrent: int = DEFAULT_MAX_CONCURRENT,
    progress_callback=None,
    cancellation_check=None,
) -> None:
    """Process the document and add translations for specified styles only

    Args:
        input_file: Path to input DOCX file
        output_file: Path to output DOCX file
        target_language: Target language for translation
        target_styles: List of paragraph styles to translate
        openai_client: OpenAI client instance
        use_cache: Whether to use translation caching
        cache_dir: Directory to store cache files
        parallel: Whether to use parallel processing
        max_concurrent: Maximum number of concurrent translation requests
        progress_callback: Optional callback function to update progress (signature: current, total, description)
        cancellation_check: Optional function that returns True if the process should be cancelled
    """
    if parallel:
        # Run the async version
        asyncio.run(
            process_document_parallel(
                input_file,
                output_file,
                target_language,
                target_styles,
                openai_client,
                use_cache,
                cache_dir,
                max_concurrent,
                progress_callback,
                cancellation_check,
            )
        )
        return

    # Original sequential implementation follows
    logger.info(f"Opening document: {input_file}")
    doc = Document(input_file)

    # Convert target_styles to set for faster lookups
    target_styles_set = set(target_styles)

    # Count elements to translate
    elements_to_translate = count_translatable_elements(doc, target_styles_set)

    logger.info(f"Translating document to {target_language}")
    logger.info(f"Targeting styles: {', '.join(target_styles)}")
    logger.info(f"Elements to translate: {elements_to_translate}")
    logger.info("Using streaming API to show translation progress")

    # Reset the token counters
    global total_tokens_received, total_cached_tokens
    total_tokens_received = 0
    total_cached_tokens = 0

    # Create a progress bar for sequential processing
    progress_bar = tqdm(
        total=elements_to_translate,
        desc=f"Translating (0 API + 0 cached = 0 total tokens)",
    )

    # Initialize translation cache if enabled
    cache = (
        TranslationCache(target_language, DEFAULT_MODEL, cache_dir)
        if use_cache
        else None
    )
    if cache:
        logger.info("Translation caching is enabled")

    processed = 0
    cached_hits = 0
    cached_tokens = 0

    try:
        # Process all paragraphs (including headings)
        for para in doc.paragraphs:
            # Check for cancellation request
            if cancellation_check and cancellation_check():
                raise InterruptedError("Translation was cancelled by the user")

            text = para.text.strip()

            # Skip empty paragraphs or non-target styles
            if not text or para.style.name not in target_styles_set:
                continue

            # Create a task ID for this paragraph
            task_id = f"Element {processed+1}/{elements_to_translate}"

            logger.debug(f"Translating {task_id}: {text[:40]}...")

            # Check if this is from cache for logging purposes
            if cache and cache.get(text):
                cached_hits += 1
                if (
                    cache.get(text).get("token_usage")
                    and "completion_tokens" in cache.get(text)["token_usage"]
                ):
                    cached_tokens += cache.get(text)["token_usage"]["completion_tokens"]

            # Use streaming translation with progress tracking
            translated_text = stream_translation(
                text, target_language, openai_client, cache, progress_bar, task_id
            )

            processed += 1
            progress_bar.update(1)

            # Update progress if callback is provided
            if progress_callback:
                eta_seconds = progress_bar.format_dict.get("eta", None)
                eta_str = (
                    f" (ETA: {time.strftime('%M:%S', time.gmtime(eta_seconds))})"
                    if eta_seconds is not None
                    else ""
                )
                progress_callback(
                    processed,
                    elements_to_translate,
                    f"Translating {task_id}: {text[:40]}...{eta_str}",
                )

            add_translation_to_paragraph(para, translated_text)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Check for cancellation request
                        if cancellation_check and cancellation_check():
                            raise InterruptedError(
                                "Translation was cancelled by the user"
                            )

                        text = para.text.strip()

                        # Skip empty paragraphs or non-target styles
                        if not text or para.style.name not in target_styles_set:
                            continue

                        # Create a task ID for this paragraph
                        task_id = f"Element {processed+1}/{elements_to_translate}"

                        logger.debug(f"Translating {task_id}: {text[:40]}...")

                        # Check if this is from cache for logging purposes
                        if cache and cache.get(text):
                            cached_hits += 1
                            if (
                                cache.get(text).get("token_usage")
                                and "completion_tokens"
                                in cache.get(text)["token_usage"]
                            ):
                                cached_tokens += cache.get(text)["token_usage"][
                                    "completion_tokens"
                                ]

                        # Use streaming translation with progress tracking
                        translated_text = stream_translation(
                            text,
                            target_language,
                            openai_client,
                            cache,
                            progress_bar,
                            task_id,
                        )

                        processed += 1
                        progress_bar.update(1)

                        # Update progress if callback is provided
                        if progress_callback:
                            progress_callback(
                                processed,
                                elements_to_translate,
                                f"Translating {task_id}: {text[:40]}...",
                            )

                        add_translation_to_paragraph(para, translated_text)

        # Close progress bar
        progress_bar.close()

        # Save cache to disk
        if cache:
            cache.save()

        logger.info(f"Saving translated document to: {output_file}")
        doc.save(output_file)

        # Add cache statistics to final message
        cache_stats = (
            f" ({cached_hits} from cache, {cached_tokens} cached tokens)"
            if use_cache
            else ""
        )
        logger.success(
            f"Document translation completed! Translated {processed} elements{cache_stats} to {target_language}"
        )
        logger.info(
            f"Total tokens used: {total_tokens_received} API + {total_cached_tokens} cached = {total_tokens_received + total_cached_tokens} total"
        )

        # Final progress update
        if progress_callback:
            progress_callback(
                elements_to_translate, elements_to_translate, "Translation completed!"
            )

    except InterruptedError as e:
        # Handle cancellation nicely - without full stack trace
        logger.warning(f"Translation was cancelled by the user: {str(e)}")
        progress_bar.close()
        raise
    except Exception as e:
        # Close progress bar on error
        progress_bar.close()

        # Log full stack trace for actual errors
        logger.exception("Error during translation")

        # Re-raise the exception to be handled by the caller
        raise e


@app.command()
def translate(
    input_file: str = typer.Argument(..., help="Input DOCX file path"),
    target_language: str = typer.Argument(
        DEFAULT_TARGET_LANGUAGE, help="Target language (e.g., 'Spanish', 'French')"
    ),
    output_file: Optional[str] = typer.Option(
        None,
        "--output",
        "-o",
        help="Output DOCX file path (defaults to 'translated_' + input_file)",
    ),
    styles: str = typer.Option(
        "Normal",
        "--styles",
        "-s",
        help="Comma-separated list of styles to translate (e.g., 'Normal,Heading 1,List Paragraph')",
    ),
    api_key: Optional[str] = typer.Option(
        None,
        "--api-key",
        help="OpenAI API key (or set OPENAI_API_KEY environment variable)",
    ),
    model: str = typer.Option(
        os.environ.get("OPENAI_MODEL", DEFAULT_MODEL),
        "--model",
        "-m",
        help="OpenAI model to use for translation",
    ),
    base_url: Optional[str] = typer.Option(
        os.environ.get("OPENAI_BASE_URL", None),
        "--base-url",
        help="OpenAI API base URL (for custom endpoints or proxies)",
    ),
    verbose: bool = typer.Option(
        False,
        "--verbose",
        "-v",
        help="Enable verbose logging",
    ),
    no_cache: bool = typer.Option(
        False,
        "--no-cache",
        help="Disable translation caching",
    ),
    clear_cache: bool = typer.Option(
        False,
        "--clear-cache",
        help="Clear the translation cache before starting",
    ),
    cache_dir: Path = typer.Option(
        DEFAULT_CACHE_DIR,
        "--cache-dir",
        help="Directory to store translation cache files",
    ),
    sequential: bool = typer.Option(
        False,
        "--sequential",
        help="Use sequential processing instead of parallel",
    ),
    max_concurrent: int = typer.Option(
        int(os.environ.get("OPENAI_MAX_CONCURRENT", DEFAULT_MAX_CONCURRENT)),
        "--max-concurrent",
        help="Maximum number of concurrent translation requests",
    ),
) -> None:
    """
    Translate a Word document to the target language, only processing text with specified styles.
    """
    # Configure logging level based on verbose flag
    if verbose:
        logger.remove()
        logger.add(
            sys.stderr,
            format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
            level="DEBUG",
            colorize=True,
        )
        logger.debug("Verbose logging enabled")

    # Update global model if specified
    global DEFAULT_MODEL
    if model != DEFAULT_MODEL:
        DEFAULT_MODEL = model
        logger.info(f"Using custom model: {model}")

    # Ensure cache directory exists
    if not no_cache:
        cache_dir = Path(cache_dir)
        cache_dir.mkdir(exist_ok=True, parents=True)
        logger.info(f"Using cache directory: {cache_dir}")

    # Handle cache clearing
    if clear_cache:
        try:
            cache_file = (
                Path(cache_dir)
                / f"cache_{target_language.lower().replace(' ', '_')}_{model.replace('-', '_')}.json"
            )
            if cache_file.exists():
                cache_file.unlink()
                logger.info(f"Cleared translation cache for {target_language}")
        except Exception as e:
            logger.warning(f"Failed to clear cache: {e}")

    # Convert input and output paths to Path objects
    input_path = Path(input_file)

    # Set default output file if not provided
    if output_file is None:
        output_path = input_path.parent / f"translated_{input_path.name}"
    else:
        output_path = Path(output_file)

    # Parse styles into a list
    target_styles = [style.strip() for style in styles.split(",")]

    try:
        # Setup OpenAI client
        client = setup_openai_client(api_key, base_url)

        # Process the document - CLI usage doesn't need cancellation
        process_document(
            input_path,
            output_path,
            target_language,
            target_styles,
            client,
            use_cache=not no_cache,
            cache_dir=cache_dir,
            parallel=not sequential,
            max_concurrent=max_concurrent,
            progress_callback=None,
            cancellation_check=None,  # CLI doesn't need cancellation
        )
    except InterruptedError as e:
        logger.warning(f"Translation was cancelled by the user: {str(e)}")
        raise typer.Exit(code=2)
    except Exception as e:
        logger.exception(f"An error occurred")
        raise typer.Exit(code=1)


@app.command()
def clear_caches(
    cache_dir: Path = typer.Option(
        os.environ.get("DOCX_TRANSLATOR_CACHE_DIR", DEFAULT_CACHE_DIR),
        "--cache-dir",
        help="Directory where cache files are stored",
    ),
) -> None:
    """Clear all translation caches"""
    try:
        cache_dir = Path(cache_dir)
        if not cache_dir.exists():
            logger.info("No cache directory found.")
            return

        count = 0
        for cache_file in cache_dir.glob("cache_*.json"):
            cache_file.unlink()
            count += 1

        logger.success(f"Cleared {count} translation cache files from {cache_dir}")
    except Exception as e:
        logger.exception(f"Error clearing caches")
        raise typer.Exit(code=1)


@app.command()
def serve(
    port: int = typer.Option(
        int(os.environ.get("STREAMLIT_PORT", 8501)),
        "--port",
        "-p",
        help="Port to run the Streamlit server on",
    ),
    cache_dir: Path = typer.Option(
        os.environ.get("DOCX_TRANSLATOR_CACHE_DIR", DEFAULT_CACHE_DIR),
        "--cache-dir",
        help="Directory to store translation cache files",
    ),
    verbose: bool = typer.Option(
        os.environ.get("DOCX_TRANSLATOR_VERBOSE", "").lower() in ("true", "1", "yes"),
        "--verbose",
        "-v",
        help="Enable verbose logging",
    ),
) -> None:
    """Start a Streamlit web interface for document translation"""
    # Configure logging level based on verbose flag
    if verbose:
        logger.remove()
        logger.add(
            sys.stderr,
            format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
            level="DEBUG",
            colorize=True,
        )
        logger.debug("Verbose logging enabled")

    import subprocess

    # Ensure the cache directory exists
    cache_dir = Path(cache_dir)
    cache_dir.mkdir(exist_ok=True, parents=True)
    logger.info(f"Using cache directory: {cache_dir}")

    logger.info(f"Starting Streamlit server on port {port}")

    # Get the directory of the current script
    script_dir = Path(__file__).parent.absolute()
    streamlit_script = script_dir / "streamlit_app.py"

    # Start the Streamlit server
    try:
        subprocess.run(
            [
                "streamlit",
                "run",
                str(streamlit_script),
                "--server.port",
                str(port),
                "--",
                "--cache-dir",
                str(cache_dir),
                "--model",
                os.environ.get("OPENAI_MODEL", DEFAULT_MODEL),
                "--target-language",
                os.environ.get("TARGET_LANGUAGE", DEFAULT_TARGET_LANGUAGE),
            ],
            check=True,
        )
    except KeyboardInterrupt:
        logger.info("Streamlit server stopped")
    except Exception as e:
        logger.exception(f"Failed to start Streamlit server: {e}")
        raise typer.Exit(code=1)


if __name__ == "__main__":
    app()
